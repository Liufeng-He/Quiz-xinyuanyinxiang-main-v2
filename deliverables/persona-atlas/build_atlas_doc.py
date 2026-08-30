from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "心院印象_新人物图鉴方案.docx"
IMG = ROOT / "images"

PURPLE = RGBColor(67, 54, 105)
TEAL = RGBColor(31, 102, 102)
GOLD = RGBColor(183, 130, 38)
INK = RGBColor(36, 31, 48)
MUTED = RGBColor(101, 94, 112)

TYPES = [
    ("AVE", "PRO-E", "社牛学霸", "亲近 · 鲜明 · 探索", "人很会社交，开口却全是方法、数据和证据。", "放大镜、数据图、开放姿态", "AVE_PRO-E.png"),
    ("AVC", "SUNNY", "人形小太阳", "亲近 · 鲜明 · 陪伴", "自带活人感，走到哪里都能把气氛带起来。", "灯笼、活动手册、明亮校园", "AVC_SUNNY.png"),
    ("ALE", "ACE-I", "静音学神", "亲近 · 潜行 · 探索", "平时保持静音，关键问题从不掉线。", "图书馆、镜片、收拢姿态", "ALE_ACE-I.png"),
    ("ALC", "WARM-I", "暖心淡人", "亲近 · 潜行 · 陪伴", "情绪稳稳的，不抢镜，但需要时一定在。", "雨夜、热饮、备用雨伞", "ALC_WARM-I.png"),
    ("WVE", "BOSS", "高配大佬", "观望 · 鲜明 · 探索", "能力与气场拉满，只是距离感也拉满。", "控制台、分析仪、冷静站姿", "WVE_BOSS.png"),
    ("WVC", "POP", "校园显眼包", "观望 · 鲜明 · 陪伴", "全校都眼熟，但和你还停留在‘见过’。", "活动展板、彩旗、窗口装置", "WVC_POP.png"),
    ("WLE", "HIDDEN", "隐藏款学神", "观望 · 潜行 · 探索", "江湖一直有传说，本人却很少刷新。", "雾中书库、棱镜、研究卷轴", "WLE_HIDDEN.png"),
    ("WLC", "LOCKED", "待解锁搭子", "观望 · 潜行 · 陪伴", "看起来可能合拍，只是认识进度还是 0%。", "半开的门、邀请信、暖色门缝", "WLC_LOCKED.png"),
]


def set_font(run, size=None, bold=None, color=INK, latin="Hiragino Sans GB", east="Hiragino Sans GB"):
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east)
    fonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_text(doc, text, size=11, bold=False, color=INK, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Hiragino Sans GB"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
normal._element.rPr.rFonts.set(qn("w:cs"), "Hiragino Sans GB")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("心院印象 · PERSONA ATLAS"), size=8.5, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("北京大学心理与认知科学学院印象调研 · 方案稿"), size=8, color=MUTED)

# Cover
add_text(doc, "XINYUAN IMPRESSION", 10, True, GOLD, 18, WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "心院印象", 31, True, PURPLE, 2, WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "新人物图鉴方案", 21, True, TEAL, 20, WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "八种现实人物原型，让受访者一眼认出心中的心院", 12.5, False, MUTED, 22, WD_ALIGN_PARAGRAPH.CENTER)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_after = Pt(20)
cover.add_run().add_picture(str(IMG / "AVE_PRO-E.png"), height=Inches(4.35))

callout = add_text(doc, "编码逻辑", 12, True, RGBColor(255,255,255), 3)
shade_paragraph(callout, "433669")
for line in [
    "A / W：亲近 / 观望（好感度）",
    "V / L：鲜明 / 潜行（存在感）",
    "E / C：探索 / 陪伴（关注朝向）",
]:
    add_text(doc, line, 10.5, False, INK, 3)

doc.add_page_break()

# Overview
add_text(doc, "01  图鉴总览", 22, True, PURPLE, 5)
add_text(doc, "命名原则：现实人物原型优先，网络表达点到为止；避免过度油腻、性别绑定和短周期梗。", 10.5, False, MUTED, 12)

table = doc.add_table(rows=1, cols=4)
table.autofit = False
widths = [Inches(0.85), Inches(1.05), Inches(1.55), Inches(3.0)]
headers = ["原编码", "新编码", "人物原型", "第一印象"]
for i, (cell, width, label) in enumerate(zip(table.rows[0].cells, widths, headers)):
    cell.width = width
    cell.vertical_alignment = 1
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(label), 9.5, True, RGBColor(255,255,255))
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "433669"); tcpr.append(shd)

for idx, (old, new, name, axes, impression, visual, filename) in enumerate(TYPES):
    cells = table.add_row().cells
    vals = [old, new, name, impression]
    for i, (cell, width, val) in enumerate(zip(cells, widths, vals)):
        cell.width = width
        cell.vertical_alignment = 1
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(val), 9.3, i in (0,1,2), PURPLE if i == 1 else INK)
        if idx % 2:
            tcpr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F4F1F7"); tcpr.append(shd)

add_text(doc, "推荐上线方式", 14, True, TEAL, 5)
add_text(doc, "结果页同时保留原编码与新编码，例如“AVE · PRO-E”，人物原型作为主标题，第一印象作为一句话副标题。", 10.5, False, INK, 5)
add_text(doc, "本方案只替换结果呈现，不改变三维计算、阈值或八类评分逻辑。", 10.5, False, MUTED, 5)

# Individual cards
for page_no, (old, new, name, axes, impression, visual, filename) in enumerate(TYPES, start=2):
    doc.add_page_break()
    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(4)
    set_font(top.add_run(f"{old}  ·  {new}"), 12, True, GOLD)
    title = add_text(doc, name, 25, True, PURPLE, 3)
    add_text(doc, axes, 11, True, TEAL, 10)

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_after = Pt(12)
    pic.add_run().add_picture(str(IMG / filename), height=Inches(5.9))

    quote = add_text(doc, f"“{impression}”", 13.5, True, INK, 10, WD_ALIGN_PARAGRAPH.CENTER)
    shade_paragraph(quote, "F4F1F7")
    add_text(doc, f"视觉锚点  {visual}", 10.5, False, MUTED, 4, WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(OUT)
